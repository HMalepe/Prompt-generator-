#!/usr/bin/env python3
"""
export_docx.py — turn a forged leaf (the 50+ prompt product) into a polished,
sellable Word document, then a PDF.

  python export_docx.py --json library/builds/<STAMP>/websites/<sub>/<leaf>.json
  python export_docx.py --json <leaf.json> --no-pdf      # Word only

Word is the source of truth; PDF is rendered via LibreOffice when available,
or Microsoft Word COM on Windows as a fallback. Output lands next to the
.json as <leaf>.docx / <leaf>.pdf.

The docx skill recommends docx-js; python-docx is used here to keep the toolchain
single-language (pip only). Output is verified by rendering to PDF and viewing.
"""
import json, argparse, subprocess, shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK   = RGBColor(0x14,0x1A,0x22)
FORGE = RGBColor(0xB4,0x6A,0x12)   # print-friendly amber
MUTE  = RGBColor(0x5A,0x66,0x74)
CODEBG= "EEF1F4"

def shade(par, hexfill):
    pPr = par._p.get_or_add_pPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"),"clear"); sh.set(qn("w:fill"),hexfill)
    pPr.append(sh)

def border(par):
    pPr = par._p.get_or_add_pPr(); pb = OxmlElement("w:pBdr")
    for side in ("top","bottom","left","right"):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"),"single"); e.set(qn("w:sz"),"4")
        e.set(qn("w:space"),"6"); e.set(qn("w:color"),"D8DEE5")
        pb.append(e)
    pPr.append(pb)

def build(data, out_docx):
    d = Document()
    n = d.sections[0]
    n.page_width, n.page_height = Inches(8.5), Inches(11)
    for m in ("top_margin","bottom_margin"): setattr(n, m, Inches(0.9))
    for m in ("left_margin","right_margin"): setattr(n, m, Inches(1.0))

    base = d.styles["Normal"].font; base.name="Calibri"; base.size=Pt(10.5)

    # cover
    d.add_paragraph().add_run().add_break()
    eyebrow = d.add_paragraph(); r = eyebrow.add_run(data["category"].replace("-"," ").upper())
    r.font.size=Pt(11); r.font.color.rgb=FORGE; r.bold=True; eyebrow.alignment=WD_ALIGN_PARAGRAPH.CENTER
    title = d.add_paragraph(); r = title.add_run(data["leaf"])
    r.font.size=Pt(30); r.bold=True; r.font.color.rgb=INK; title.alignment=WD_ALIGN_PARAGRAPH.CENTER
    sub = d.add_paragraph(); r = sub.add_run(data.get("one_liner","") or "")
    r.font.size=Pt(12); r.italic=True; r.font.color.rgb=MUTE; sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
    meta = d.add_paragraph(); r = meta.add_run(
        f'{data["prompt_count"]} step-by-step prompts   ·   {data["subcategory"]}   ·   '
        f'{data.get("difficulty","").title()}')
    r.font.size=Pt(9.5); r.font.color.rgb=MUTE; meta.alignment=WD_ALIGN_PARAGRAPH.CENTER

    howto = d.add_paragraph()
    r = howto.add_run("How to use: paste each prompt, in order, into your AI tool. "
                      "Each step builds on the last. Replace anything in [SQUARE BRACKETS] with your details.")
    r.font.size=Pt(9.5); r.font.color.rgb=MUTE; howto.alignment=WD_ALIGN_PARAGRAPH.CENTER
    howto.paragraph_format.space_before=Pt(14)
    d.add_paragraph().add_run().add_break()

    # prompts
    for p in data["prompts"]:
        h = d.add_paragraph(); h.paragraph_format.space_before=Pt(16); h.paragraph_format.keep_with_next=True
        r = h.add_run(f'STEP {p["step"]:02d}'); r.bold=True; r.font.color.rgb=FORGE; r.font.size=Pt(11)
        r2 = h.add_run(f'   {p.get("goal","")}'); r2.bold=True; r2.font.color.rgb=INK; r2.font.size=Pt(12)

        lbl = d.add_paragraph(); r=lbl.add_run("PASTE THIS"); r.font.size=Pt(8); r.bold=True; r.font.color.rgb=MUTE
        lbl.paragraph_format.space_before=Pt(4); lbl.paragraph_format.space_after=Pt(2)

        code = d.add_paragraph(); shade(code, CODEBG); border(code)
        code.paragraph_format.space_after=Pt(4)
        rr = code.add_run(p.get("prompt","")); rr.font.name="Consolas"; rr.font.size=Pt(9.5)
        rr.font.color.rgb=INK

        fi = (p.get("fill_in") or "").strip()
        if fi and fi.lower()!="nothing":
            f = d.add_paragraph(); r=f.add_run("Fill in: "); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=MUTE
            r2=f.add_run(fi); r2.font.size=Pt(9); r2.font.color.rgb=FORGE
        ex = (p.get("expected") or "").strip()
        if ex:
            e = d.add_paragraph(); r=e.add_run("Expected: "); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=MUTE
            r2=e.add_run(ex); r2.font.size=Pt(9); r2.font.color.rgb=MUTE

    d.save(out_docx)
    return out_docx

def to_pdf_word(docx_path):
    """Convert via Microsoft Word COM (Windows). No pywin32 required."""
    import sys
    if sys.platform != "win32":
        return None
    docx = Path(docx_path).resolve()
    pdf = docx.with_suffix(".pdf")
    ps = (
        f'$docx = "{docx}"; $pdf = "{pdf}"; '
        "$w = New-Object -ComObject Word.Application; $w.Visible = $false; "
        "try { $d = $w.Documents.Open($docx); $d.SaveAs([ref]$pdf, [ref]17); $d.Close() } "
        "finally { $w.Quit() }"
    )
    r = subprocess.run(
        ["powershell", "-NoProfile", "-Command", ps],
        capture_output=True, text=True,
    )
    return pdf if r.returncode == 0 and pdf.exists() else None

def to_pdf(docx_path):
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir",
                        str(Path(docx_path).parent), str(docx_path)], check=True,
                       stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        return Path(docx_path).with_suffix(".pdf")
    pdf = to_pdf_word(docx_path)
    if pdf:
        return pdf
    print("No PDF converter found — install LibreOffice or Microsoft Word, or use --no-pdf.")
    return None

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--json", required=True)
    ap.add_argument("--no-pdf", action="store_true")
    a = ap.parse_args()
    data = json.loads(Path(a.json).read_text())
    out = Path(a.json).with_suffix(".docx")
    build(data, str(out)); print("Word:", out)
    if not a.no_pdf:
        pdf = to_pdf(str(out))
        if pdf: print("PDF :", pdf)

if __name__ == "__main__":
    main()
